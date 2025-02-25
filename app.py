import streamlit as st
import os
import subprocess
import sys
import fitz  # PyMuPDF for PDFs
import docx
import pandas as pd
import chromadb
from llama_index.core import SimpleDirectoryReader, Document, VectorStoreIndex, StorageContext
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.llms.ollama import Ollama
from llama_index.core.node_parser import SimpleNodeParser
from chromadb.config import Settings
from streamlit_pdf_viewer import pdf_viewer

# Function to read files
def read_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text = []
        for page in doc:
            text.append(page.get_text("text"))  # Extract text mode
        return "\n".join(text).strip() if text else "No extractable text found."
    except Exception as e:
        return f"Error reading PDF: {e}"

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text.strip())  # Remove extra spaces
        return "\n".join(text)
    except Exception as e:
        return f"Error reading DOCX: {e}"

def read_excel(file_path):
    try:
        df = pd.read_excel(file_path, engine="openpyxl")
        return df.to_string(index=False)  # Remove index for cleaner output
    except Exception as e:
        return f"Error reading XLSX: {e}"

# Streamlit UI Layout
st.set_page_config(layout="wide")
col1, col2 = st.columns([1, 2])

with col1:
    st.title("Document Viewer")
    uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, XLSX)", type=["pdf", "docx", "xlsx"])
    
    if uploaded_file is not None:
        temp_dir = "temp"
        os.makedirs(temp_dir, exist_ok=True)
        file_path = os.path.join(temp_dir, uploaded_file.name)
        
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        if uploaded_file.name.endswith(".pdf"):
            # Display the PDF
            pdf_viewer(file_path)
            document_text = read_pdf(file_path)
        elif uploaded_file.name.endswith(".docx"):
            document_text = read_docx(file_path)
        elif uploaded_file.name.endswith(".xlsx"):
            document_text = read_excel(file_path)
        
        st.session_state["document_text"] = document_text
        st.text_area("Extracted Text:", document_text, height=600)

# Right column for chat
with col2:
    st.title("Chat with Document")
    embed_model = HuggingFaceEmbedding()
    DB_PATH = "./huyen_chromadb"
    db = chromadb.PersistentClient(path=DB_PATH, settings=Settings(allow_reset=True))
    chroma_collection = db.get_or_create_collection('huyen_collection')
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    
    parser = SimpleNodeParser.from_defaults(chunk_size=256)  # Adjusted chunk size
    index = None
    try:
        if "document_text" in st.session_state:
            documents = [Document(text=st.session_state["document_text"], metadata={"source": uploaded_file.name})]
            index = VectorStoreIndex.from_documents(documents, storage_context=storage_context, embed_model=embed_model, llm=None)
        else:
            # Handle the case when no document is uploaded
            st.info("Please upload a document to start chatting.")
    except Exception as e:
        st.error(f"Error creating index: {e}")
    
    def generate_response(input_text):
        try:
            llm = Ollama(model='deepseek-r1:7b', request_timeout=128.0, base_url='http://localhost:11434')
            query_engine = index.as_query_engine(llm=llm, similarity_top_k=4)  # Adjusted similarity threshold
            response = query_engine.query(input_text)
            return response
        except Exception as e:
            return f"Error processing query: {e}"
    
    if "document_text" in st.session_state:
        if prompt := st.text_input("Ask something about the document..."):
            st.write(f"User: {prompt}")
            response = generate_response(prompt)
            st.write(f"Assistant: {response}")
    
    if 'chat_history' not in st.session_state:
        st.session_state['chat_history'] = []
